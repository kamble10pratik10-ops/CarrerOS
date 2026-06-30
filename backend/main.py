import os
import uuid
import shutil
from fastapi import FastAPI, UploadFile, File, HTTPException, status, Depends, Request, WebSocket, WebSocketDisconnect, Query
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from passlib.context import CryptContext

# Load environment variables from .env
load_dotenv(override=True)

from src.lib.lemma.datastore import (
    get_applications,
    get_weekly_velocity,
    get_active_nudges,
    update_application,
    delete_application,
    save_profile,
    get_profile,
    get_discoverable_profiles,
    follow_user,
    unfollow_user,
    search_profiles,
    get_profiles_by_emails,
    save_chat_message,
    get_chat_history,
    get_conversations,
    get_companies,
    save_company_interest,
    get_user_company_interests
)
from src.lib.lemma.workflows import trigger_application_workflow
from src.services.ai_service import parse_file_with_gemini, chat_with_agent, chat_with_recruiter_agent, evaluate_interview_performance, generate_interview_questions, chat_with_mentor, generate_network_matches, generate_ai_introduction, generate_career_twin, analyze_goal_performance, generate_goal_recommendations
from src.lib.lemma.auth import create_access_token, verify_token
from src.lib.lemma.auth_store import user_exists, get_user_by_email, create_user, update_password
from src.lib.lemma.chat_manager import chat_manager
from src.lib.lemma.coding_store import save_coding_session, get_coding_sessions, update_coding_session, get_coding_progress, update_coding_progress
from src.services.coding_mentor_service import (
    generate_coding_hint, review_code_live, debug_code, analyze_complexity,
    simulate_execution, generate_reflection, mentor_chat as coding_mentor_chat,
    estimate_user_level
)

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

app = FastAPI(title="CareerOS API", version="1.0.0")

import os
os.makedirs("uploads/profiles", exist_ok=True)
os.makedirs("uploads/chat", exist_ok=True)
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

# CORS middleware configuration matching Node.js backend

origins = [
    "http://localhost:3000",
    "https://carrer-os-sand.vercel.app",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class RegisterRequest(BaseModel):
    email: str
    mobile: str
    password: str
    recoveryQuestion: str | None = None
    recoveryAnswer: str | None = None

class LoginRequest(BaseModel):
    email: str
    password: str

class AnalyzeRequest(BaseModel):
    jdText: str
    resumeText: str
    companyName: str | None = None
    roleName: str | None = None

class UpdateRequest(BaseModel):
    id: str | None = None
    updates: dict | None = None
    profile: dict | None = None

class ChatRequest(BaseModel):
    message: str
    chatHistory: list | None = []
    resumeText: str | None = ""
    currentJdText: str | None = ""

class EvaluateRequest(BaseModel):
    chatHistory: list | None = []
    currentJdText: str | None = ""

class MentorChatRequest(BaseModel):
    message: str
    chatHistory: list | None = []
    resumeText: str | None = ""
    profile: dict | None = {}
    attachments: list | None = []

class GenerateQuestionsRequest(BaseModel):
    targetRole: str | None = ""
    resumeText: str | None = ""
    skills: str | None = ""
    questionTypes: list | None = ["technical", "behavioral", "situational"]

class NetworkIntroRequest(BaseModel):
    targetEmail: str

class ChatSendRequest(BaseModel):
    receiver: str
    text: str
    mediaUrls: list | None = []

# --- Auth Helper ---
def get_current_user(request: Request):
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Not authenticated")
    token = auth_header.split(" ")[1]
    email = verify_token(token)
    if not email:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    return email

# --- Auth Routes ---

@app.post("/api/auth/register")
async def register(req: RegisterRequest):
    if not req.email or not req.password or not req.mobile:
        raise HTTPException(status_code=400, detail="Email, mobile and password are required")
    if user_exists(req.email):
        raise HTTPException(status_code=409, detail="An account with this email already exists")
    password_hash = pwd_context.hash(req.password)
    recovery_answer_hash = pwd_context.hash(req.recoveryAnswer) if req.recoveryAnswer else None
    create_user(req.email, req.mobile, password_hash, req.recoveryQuestion, recovery_answer_hash)
    token = create_access_token(req.email)
    return {"success": True, "token": token, "email": req.email.lower()}

@app.post("/api/auth/login")
async def login(req: LoginRequest):
    user = get_user_by_email(req.email)
    if not user or not pwd_context.verify(req.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token = create_access_token(req.email)
    return {"success": True, "token": token, "email": user["email"]}

class ForgotPasswordVerifyRequest(BaseModel):
    email: str

class ForgotPasswordResetRequest(BaseModel):
    email: str
    recoveryAnswer: str
    newPassword: str

@app.post("/api/auth/forgot-password/verify-email")
async def forgot_password_verify_email(req: ForgotPasswordVerifyRequest):
    user = get_user_by_email(req.email)
    if not user:
        raise HTTPException(status_code=404, detail="Account not found")
    if not user.get("recovery_question"):
        raise HTTPException(status_code=400, detail="No recovery question set for this account")
    return {"success": True, "recoveryQuestion": user["recovery_question"]}

@app.post("/api/auth/forgot-password/reset")
async def forgot_password_reset(req: ForgotPasswordResetRequest):
    user = get_user_by_email(req.email)
    if not user:
        raise HTTPException(status_code=404, detail="Account not found")
    if not user.get("recovery_answer_hash") or not pwd_context.verify(req.recoveryAnswer, user["recovery_answer_hash"]):
        raise HTTPException(status_code=401, detail="Incorrect recovery answer")
    
    new_password_hash = pwd_context.hash(req.newPassword)
    success = update_password(req.email, new_password_hash)
    if not success:
        raise HTTPException(status_code=500, detail="Failed to update password")
    
    return {"success": True, "message": "Password updated successfully"}

# --- Profile Routes ---

@app.get("/api/profile")
async def get_profile_route(email: str = Depends(get_current_user)):
    profile = get_profile(email)
    return {"success": True, "profile": profile}

@app.put("/api/profile")
async def update_profile_route(req: UpdateRequest, email: str = Depends(get_current_user)):
    existing = get_profile(email) or {}
    if req.profile:
        existing.update(req.profile)
    saved = save_profile(email, existing)
    return {"success": True, "profile": saved}

# 1. File Upload / Parser Route (Gemini OCR & PDF Parsing)
@app.post("/api/parse-file")
async def parse_file(file: UploadFile = File(...)):
    try:
        file_bytes = await file.read()
        mime_type = file.content_type or ""
        file_name = file.filename or "unknown"
        
        print(f"[Backend] Processing file upload: {file_name} ({mime_type})")
        extracted_text = parse_file_with_gemini(file_bytes, mime_type, file_name)
        
        return {
            "success": True,
            "text": extracted_text,
            "fileName": file_name,
            "mimeType": mime_type
        }
    except Exception as e:
        print(f"Error in parse-file route: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e) or "Failed to parse file"
        )

# 2. Application Analysis Workflow Route (Groq + Gemini Triage Scorer)
@app.post("/api/analyze")
async def analyze_application(req: AnalyzeRequest, email: str = Depends(get_current_user)):
    try:
        print(f"[Backend] Triggering analysis workflow for {email}...")
        saved_app = await trigger_application_workflow(email, {
            "jdText": req.jdText,
            "resumeText": req.resumeText,
            "companyName": req.companyName,
            "roleName": req.roleName
        })
        
        weekly_velocity = get_weekly_velocity(email)
        active_nudges = get_active_nudges(email)
        all_apps = get_applications(email)
        
        return {
            "success": True,
            "application": saved_app,
            "weeklyVelocity": weekly_velocity,
            "activeNudges": active_nudges,
            "allApps": all_apps
        }
    except Exception as e:
        print(f"Error in analyze route: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e) or "Failed to analyze application"
        )

# 3. Get Application pipeline board and stats
@app.get("/api/analyze")
async def get_analysis_dashboard(email: str = Depends(get_current_user)):
    try:
        all_apps = get_applications(email)
        weekly_velocity = get_weekly_velocity(email)
        active_nudges = get_active_nudges(email)
        
        return {
            "success": True,
            "applications": all_apps,
            "weeklyVelocity": weekly_velocity,
            "activeNudges": active_nudges
        }
    except Exception as e:
        print(f"Error fetching applications: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to fetch application states"
        )

class CareerTwinRequest(BaseModel):
    resumeText: str = ""
    targetRole: str = ""
    skills: str = ""
    projects: str = ""
    experience: str = ""
    education: str = ""

@app.post("/api/career-twin/analyze")
async def analyze_career_twin(req: CareerTwinRequest, email: str = Depends(get_current_user)):
    try:
        print(f"[Backend] Generating Career Twin for {email}...")
        result = await generate_career_twin(req.dict())
        return {"success": True, "data": result}
    except Exception as e:
        print(f"Error generating Career Twin: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e) or "Failed to generate Career Twin"
        )

class GoalAnalyzeRequest(BaseModel):
    goal: dict
    history: dict

@app.post("/api/goals/analyze")
async def analyze_goal_endpoint(req: GoalAnalyzeRequest, email: str = Depends(get_current_user)):
    try:
        print(f"[Backend] Analyzing goal performance for {email}...")
        result = await analyze_goal_performance(req.dict())
        return {"success": True, "data": result}
    except Exception as e:
        print(f"Error analyzing goal: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e) or "Failed to analyze goal"
        )

class GoalRecommendRequest(BaseModel):
    targetRole: str = ""
    skills: str = ""
    missingSkills: list = []

@app.post("/api/goals/recommend")
async def recommend_goals_endpoint(req: GoalRecommendRequest, email: str = Depends(get_current_user)):
    try:
        print(f"[Backend] Recommending goals for {email}...")
        result = await generate_goal_recommendations(req.dict())
        return {"success": True, "data": result}
    except Exception as e:
        print(f"Error recommending goals: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e) or "Failed to recommend goals"
        )

# 4. Update application status or user profile
@app.put("/api/applications")
async def update_applications_or_profile(req: UpdateRequest, email: str = Depends(get_current_user)):
    try:
        if req.profile is not None:
            print(f"[Backend] Updating user profile for {email}...")
            saved = save_profile(email, req.profile)
            return {"success": True, "profile": saved}
            
        if not req.id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Application ID is required"
            )
            
        print(f"[Backend] Updating application status: {req.id} for {email}")
        updated_app = update_application(email, req.id, req.updates or {})
        
        if not updated_app:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Application not found"
            )
            
        return {"success": True, "application": updated_app}
    except HTTPException as he:
        raise he
    except Exception as e:
        print(f"Error in applications PUT route: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e) or "Failed to update"
        )

# 5. Delete an application
@app.delete("/api/applications")
async def delete_app(id: str, email: str = Depends(get_current_user)):
    try:
        print(f"[Backend] Deleting application {id} for {email}")
        success = delete_application(email, id)
        return {"success": success}
    except Exception as e:
        print(f"Error in applications DELETE route: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to delete application"
        )

# 5.b Download Tailored Resume
import docx
import json
import tempfile
from starlette.background import BackgroundTask

@app.get("/api/applications/{app_id}/download-resume")
async def download_resume(app_id: str, email: str = Depends(get_current_user)):
    try:
        all_apps = get_applications(email)
        app_data = next((a for a in all_apps if a.get("id") == app_id), None)
        if not app_data:
            raise HTTPException(status_code=404, detail="Application not found")
            
        profile = get_profile(email)
        resume_text = profile.get("resumeText", "") if profile else ""
        if not resume_text:
            raise HTTPException(status_code=400, detail="No base resume found")
            
        # Parse tailored bullets
        tailored_bullets_raw = app_data.get("tailoredBullets")
        if not tailored_bullets_raw:
            tailored_bullets = []
        else:
            try:
                tailored_bullets = json.loads(tailored_bullets_raw) if isinstance(tailored_bullets_raw, str) else tailored_bullets_raw
            except:
                tailored_bullets = []
            
        if not isinstance(tailored_bullets, list):
            tailored_bullets = []
            
        # Reconstruct resume text by replacing original with tailored
        final_resume = resume_text
        for bullet in tailored_bullets:
            original = bullet.get("original", "").strip()
            tailored = bullet.get("tailored", "").strip()
            if original and tailored and original in final_resume:
                final_resume = final_resume.replace(original, tailored)
                
        # Create a DOCX
        doc = docx.Document()
        
        # Add basic formatting
        for paragraph in final_resume.split('\n'):
            p = paragraph.strip()
            if p:
                if p.isupper() and len(p) < 40:
                    doc.add_heading(p, level=2)
                elif p.startswith(('-', '*')):
                    doc.add_paragraph(p[1:].strip(), style='List Bullet')
                else:
                    doc.add_paragraph(p)
                    
        # Save to temp file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        
        return FileResponse(
            temp_file.name, 
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"Tailored_Resume_{app_data.get('company', 'Company')}.docx",
            background=BackgroundTask(lambda p: os.unlink(p) if os.path.exists(p) else None, temp_file.name)
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error downloading resume: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to generate document"
        )

# 6. Conversational Chat Agent (Career Concierge)
@app.post("/api/chat")
async def chat(req: ChatRequest, email: str = Depends(get_current_user)):
    try:
        print(f"[Backend] Forwarding message to Career Concierge for {email}...")
        reply = await chat_with_agent({
            "message": req.message,
            "chatHistory": req.chatHistory or [],
            "resumeText": req.resumeText,
            "currentJdText": req.currentJdText
        })
        
        return {
            "success": True,
            "reply": reply
        }
    except Exception as e:
        print(f"Error in chat API route: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e) or "Failed to get response from AI chat"
        )

# 7. Recruiter Simulator Chat Agent
@app.post("/api/interview-chat")
async def interview_chat(req: ChatRequest, email: str = Depends(get_current_user)):
    try:
        print(f"[Backend] Forwarding message to Recruiter Simulator for {email}...")
        reply = await chat_with_recruiter_agent({
            "message": req.message,
            "chatHistory": req.chatHistory or [],
            "resumeText": req.resumeText,
            "currentJdText": req.currentJdText
        })
        
        return {
            "success": True,
            "reply": reply
        }
    except Exception as e:
        print(f"Error in interview chat API route: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e) or "Failed to get response from Recruiter Simulator"
        )

# 8. Recruiter Simulator Interview Evaluator
@app.post("/api/evaluate-interview")
async def evaluate_interview(req: EvaluateRequest, email: str = Depends(get_current_user)):
    try:
        print(f"[Backend] Forwarding interview to Evaluator for {email}...")
        evaluation = await evaluate_interview_performance({
            "chatHistory": req.chatHistory or [],
            "currentJdText": req.currentJdText
        })
        
        return {
            "success": True,
            "evaluation": evaluation
        }
    except Exception as e:
        print(f"Error in evaluate-interview API route: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e) or "Failed to evaluate interview performance"
        )

# 9. AI Mentor Chat
@app.post("/api/mentor/chat")
async def mentor_chat(req: MentorChatRequest, email: str = Depends(get_current_user)):
    try:
        print(f"[Backend] Forwarding message to AI Mentor for {email}...")
        reply = await chat_with_mentor({
            "message": req.message,
            "chatHistory": req.chatHistory or [],
            "resumeText": req.resumeText,
            "profile": req.profile or {},
            "attachments": req.attachments or []
        })
        return {"success": True, "reply": reply}
    except Exception as e:
        print(f"Error in mentor chat API route: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e) or "Failed to get response from AI Mentor"
        )

# 10. Interview Question Generator
@app.post("/api/generate-questions")
async def generate_questions(req: GenerateQuestionsRequest, email: str = Depends(get_current_user)):
    try:
        print(f"[Backend] Generating interview questions for {email}...")
        result = await generate_interview_questions({
            "targetRole": req.targetRole or "",
            "resumeText": req.resumeText or "",
            "skills": req.skills or "",
            "questionTypes": req.questionTypes or ["technical", "behavioral", "situational"]
        })
        return {"success": True, **result}
    except Exception as e:
        print(f"Error in generate-questions API route: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e) or "Failed to generate interview questions"
        )

# 11. Career Connect Matchmaking
@app.get("/api/network/matches")
async def get_network_matches(email: str = Depends(get_current_user)):
    try:
        current_profile = get_profile(email) or {}
        if not current_profile.get("discoverable"):
            return {"success": True, "matches": [], "message": "You must opt-in to discoverability to see matches."}
            
        candidate_profiles = get_discoverable_profiles(exclude_email=email)
        if not candidate_profiles:
            return {"success": True, "matches": []}
            
        print(f"[Backend] Generating network matches for {email} out of {len(candidate_profiles)} candidates...")
        result = await generate_network_matches({
            "currentProfile": current_profile,
            "candidateProfiles": candidate_profiles
        })
        return {"success": True, "matches": result.get("matches", [])}
    except Exception as e:
        print(f"Error in network matches API route: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate network matches")

@app.post("/api/network/intro")
async def generate_network_intro(req: NetworkIntroRequest, email: str = Depends(get_current_user)):
    try:
        current_profile = get_profile(email) or {}
        target_profile = get_profile(req.targetEmail) or {}
        
        if not target_profile:
            raise HTTPException(status_code=404, detail="Target profile not found")
            
        print(f"[Backend] Generating intro from {email} to {req.targetEmail}...")
        intro_text = await generate_ai_introduction({
            "currentProfile": current_profile,
            "targetProfile": target_profile
        })
        return {"success": True, "intro": intro_text}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error in network intro API route: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate intro")

class FollowRequest(BaseModel):
    targetEmail: str

@app.post("/api/profile/photo")
async def upload_profile_photo(file: UploadFile = File(...), email: str = Depends(get_current_user)):
    try:
        ext = file.filename.split(".")[-1]
        filename = f"{uuid.uuid4().hex}.{ext}"
        filepath = os.path.join("uploads", "profiles", filename)
        
        with open(filepath, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        photo_url = f"/uploads/profiles/{filename}"
        
        profile = get_profile(email) or {}
        profile["photoUrl"] = photo_url
        save_profile(email, profile)
        
        return {"success": True, "photoUrl": photo_url}
    except Exception as e:
        print(f"Error uploading photo: {e}")
        raise HTTPException(status_code=500, detail="Failed to upload photo")

@app.get("/api/network/search")
async def network_search(q: str, email: str = Depends(get_current_user)):
    try:
        profiles = search_profiles(q, exclude_email=email)
        return {"success": True, "profiles": profiles}
    except Exception as e:
        raise HTTPException(status_code=500, detail="Search failed")

@app.post("/api/network/follow")
async def network_follow(req: FollowRequest, email: str = Depends(get_current_user)):
    try:
        success = follow_user(email, req.targetEmail)
        return {"success": success}
    except Exception as e:
        raise HTTPException(status_code=500, detail="Follow failed")

@app.post("/api/network/unfollow")
async def network_unfollow(req: FollowRequest, email: str = Depends(get_current_user)):
    try:
        success = unfollow_user(email, req.targetEmail)
        return {"success": success}
    except Exception as e:
        raise HTTPException(status_code=500, detail="Unfollow failed")

@app.get("/api/network/users")
async def get_network_users(emails: str, email: str = Depends(get_current_user)):
    try:
        email_list = [e.strip() for e in emails.split(",") if e.strip()]
        profiles = get_profiles_by_emails(email_list)
        return {"success": True, "profiles": profiles}
    except Exception as e:
        raise HTTPException(status_code=500, detail="Failed to fetch users")

# --- Chat Routes ---

@app.post("/api/chat/send")
async def send_chat_message(req: ChatSendRequest, email: str = Depends(get_current_user)):
    try:
        msg = save_chat_message(email, req.receiver, req.text, req.mediaUrls or [])
        if msg:
            # Broadcast via WebSocket if connected
            payload = {"type": "message", "message": msg}
            await chat_manager.send_to_user(email, payload)
            await chat_manager.send_to_user(req.receiver, payload)
            return {"success": True, "message": msg}
        raise HTTPException(status_code=500, detail="Failed to save message")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/chat/history")
async def chat_history(user2: str, email: str = Depends(get_current_user)):
    try:
        history = get_chat_history(email, user2)
        return {"success": True, "messages": history}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/chat/conversations")
async def chat_conversations(email: str = Depends(get_current_user)):
    try:
        convos = get_conversations(email)
        # Enrich with profile info
        enriched = []
        for conv in convos:
            profile = get_profile(conv["email"]) or {}
            conv["name"] = profile.get("name", conv["email"])
            conv["photoUrl"] = profile.get("photoUrl", "")
            enriched.append(conv)
        return {"success": True, "conversations": enriched}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- WebSocket Chat ---

async def handle_websocket(websocket: WebSocket, token: str):
    email = verify_token(token)
    if not email:
        await websocket.close(code=4001)
        return
    await chat_manager.connect(email, websocket)
    try:
        while True:
            data = await websocket.receive_json()
            msg_type = data.get("type")
            if msg_type == "send_message":
                receiver = data.get("receiver", "").strip()
                text = data.get("text", "").strip()
                media_urls = data.get("mediaUrls") or []
                if receiver and (text or media_urls):
                    await chat_manager.handle_message(email, receiver, text, media_urls)
            elif msg_type == "typing":
                receiver = data.get("receiver", "").strip()
                if receiver:
                    await chat_manager.handle_typing(email, receiver)
    except WebSocketDisconnect:
        chat_manager.disconnect(email, websocket)
    except Exception:
        chat_manager.disconnect(email, websocket)

@app.websocket("/ws")
async def ws_endpoint(websocket: WebSocket, token: str = Query(...)):
    await handle_websocket(websocket, token)

@app.websocket("/api/chat/ws")
async def api_chat_ws_endpoint(websocket: WebSocket, token: str = Query(...)):
    await handle_websocket(websocket, token)

# --- Company Routes ---

@app.get("/api/companies")
async def list_companies(q: str = "", email: str = Depends(get_current_user)):
    try:
        companies = get_companies(q)
        interested_ids = get_user_company_interests(email)
        for c in companies:
            c["interested"] = c.get("id") in interested_ids
        return {"success": True, "companies": companies}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/companies/connect")
async def connect_company(req: NetworkIntroRequest, email: str = Depends(get_current_user)):
    try:
        ok = save_company_interest(email, req.targetEmail)
        return {"success": ok}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# --- Coding Mentor Routes ---

class CodingHintRequest(BaseModel):
    problem: dict
    code: str
    language: str
    hintLevel: int = 1
    chatHistory: list | None = []
    userLevel: str = "Beginner"

class CodingReviewRequest(BaseModel):
    problem: dict
    code: str
    language: str
    userLevel: str = "Beginner"

class CodingDebugRequest(BaseModel):
    problem: dict
    code: str
    language: str
    errorOutput: str
    errorType: str = "runtime"

class CodingComplexityRequest(BaseModel):
    problem: dict
    code: str
    language: str

class CodingRunRequest(BaseModel):
    code: str
    language: str
    testCases: list

class CodingSubmitRequest(BaseModel):
    problem: dict
    code: str
    language: str
    testCases: list
    hintsUsed: int = 0
    timeSpent: int = 0
    userLevel: str = "Beginner"

class CodingChatRequest(BaseModel):
    problem: dict
    code: str
    language: str
    message: str
    chatHistory: list | None = []
    userLevel: str = "Beginner"
    interviewMode: bool = False

class CodingProgressRequest(BaseModel):
    sessionData: dict | None = None
    progressData: dict | None = None

@app.post("/api/coding-mentor/hint")
async def coding_hint(req: CodingHintRequest, email: str = Depends(get_current_user)):
    try:
        result = await generate_coding_hint(req.problem, req.code, req.language, req.hintLevel, req.chatHistory or [], req.userLevel)
        return {"success": True, "data": result}
    except Exception as e:
        print(f"Error in coding hint: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/coding-mentor/review")
async def coding_review(req: CodingReviewRequest, email: str = Depends(get_current_user)):
    try:
        result = await review_code_live(req.problem, req.code, req.language, req.userLevel)
        return {"success": True, "data": result}
    except Exception as e:
        print(f"Error in coding review: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/coding-mentor/debug")
async def coding_debug(req: CodingDebugRequest, email: str = Depends(get_current_user)):
    try:
        result = await debug_code(req.problem, req.code, req.language, req.errorOutput, req.errorType)
        return {"success": True, "data": result}
    except Exception as e:
        print(f"Error in coding debug: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/coding-mentor/complexity")
async def coding_complexity(req: CodingComplexityRequest, email: str = Depends(get_current_user)):
    try:
        result = await analyze_complexity(req.problem, req.code, req.language)
        return {"success": True, "data": result}
    except Exception as e:
        print(f"Error in coding complexity: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/coding-mentor/run")
async def coding_run(req: CodingRunRequest, email: str = Depends(get_current_user)):
    try:
        result = await simulate_execution(req.code, req.language, req.testCases)
        return {"success": True, "data": result}
    except Exception as e:
        print(f"Error in coding run: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/coding-mentor/submit")
async def coding_submit(req: CodingSubmitRequest, email: str = Depends(get_current_user)):
    try:
        exec_result = await simulate_execution(req.code, req.language, req.testCases)
        reflection = await generate_reflection(req.problem, req.code, req.language, exec_result, req.hintsUsed, req.timeSpent, req.userLevel)
        # Save session
        session_data = {
            "problem_id": req.problem.get("id", "custom"),
            "language": req.language,
            "code": req.code,
            "status": "completed",
            "hints_used": req.hintsUsed,
            "time_spent_seconds": req.timeSpent,
            "test_results": exec_result,
            "reflection": reflection
        }
        save_coding_session(email, session_data)
        # Update progress
        progress = get_coding_progress(email) or {}
        progress["problems_solved"] = progress.get("problems_solved", 0) + (1 if all(r.get("passed") for r in exec_result.get("results", [])) else 0)
        progress["total_attempts"] = progress.get("total_attempts", 0) + 1
        progress["total_hints_used"] = progress.get("total_hints_used", 0) + req.hintsUsed
        update_coding_progress(email, progress)
        return {"success": True, "execution": exec_result, "reflection": reflection}
    except Exception as e:
        print(f"Error in coding submit: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/coding-mentor/chat")
async def coding_chat(req: CodingChatRequest, email: str = Depends(get_current_user)):
    try:
        result = await coding_mentor_chat(req.problem, req.code, req.language, req.message, req.chatHistory or [], req.userLevel, req.interviewMode)
        return {"success": True, "reply": result}
    except Exception as e:
        print(f"Error in coding chat: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/coding-mentor/progress")
async def get_coding_progress_route(email: str = Depends(get_current_user)):
    try:
        progress = get_coding_progress(email)
        sessions = get_coding_sessions(email)
        level = await estimate_user_level(progress or {})
        return {"success": True, "progress": progress, "sessions": sessions, "level": level}
    except Exception as e:
        print(f"Error getting coding progress: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/coding-mentor/progress")
async def save_coding_progress_route(req: CodingProgressRequest, email: str = Depends(get_current_user)):
    try:
        if req.sessionData:
            save_coding_session(email, req.sessionData)
        if req.progressData:
            update_coding_progress(email, req.progressData)
        return {"success": True}
    except Exception as e:
        print(f"Error saving coding progress: {e}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    # Use port 3001 to match Next.js frontend expectance
    uvicorn.run("main:app", host="127.0.0.1", port=3001, reload=False)
