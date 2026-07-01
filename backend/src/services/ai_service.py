import os
import json
import base64
from groq import Groq
import google.generativeai as genai
from src.services.lemma_orchestrator import run_mentor_workflow, run_resume_analysis_workflow

def get_groq_client():
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        return None
    try:
        return Groq(api_key=api_key)
    except Exception:
        return None

def get_gemini_client():
    api_key = os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        return None
    genai.configure(api_key=api_key)
    return genai

def clean_json_string(text):
    text = text.strip()
    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    return text.strip()

def parse_file_with_gemini(file_bytes, mime_type, file_name):
    file_ext = (file_name.rsplit('.', 1)[-1] if '.' in file_name else '').lower()
    
    # --- Plain text files: just decode ---
    if mime_type.startswith("text/") or file_ext in ("txt", "text", "csv", "md"):
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1")
    
    # --- DOCX: extract with python-docx (Gemini can't handle DOCX inline) ---
    if file_ext in ("docx",) or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            extracted = "\n".join(paragraphs)
            if extracted.strip():
                return extracted
            # If paragraphs are empty, try tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        table_text.append(" | ".join(row_data))
            return "\n".join(table_text) if table_text else "Could not extract text from this DOCX file."
        except ImportError:
            raise ValueError("python-docx is not installed. Run: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file: {e}")
    
    # --- Old .doc format: try textract-like fallback ---
    if file_ext == "doc" or mime_type == "application/msword":
        # .doc is a legacy binary format ÔÇö try to decode as text, or advise to convert
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
            # Strip binary noise
            clean = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
            if len(clean.strip()) > 50:
                return clean.strip()
        except Exception:
            pass
        raise ValueError("Old .doc format detected. Please save as .docx or PDF and try again.")
    
    # --- Images and PDFs: use Gemini Vision/Document AI ---
    gen_ai = get_gemini_client()
    if not gen_ai:
        print("[ai_service] WARNING: Gemini API key not set. Returning mock parsed text.")
        return "This is mock parsed text. Please configure GOOGLE_API_KEY to parse files."
        
    model = gen_ai.GenerativeModel('gemini-1.5-flash')
    
    part = {
        "inline_data": {
            "mime_type": mime_type,
            "data": base64.b64encode(file_bytes).decode('utf-8')
        }
    }
    
    if mime_type.startswith("image/"):
        prompt = "Extract ALL text from this image. This may be a resume or a job description screenshot. Provide the complete parsed text clearly, preserving structure."
    elif mime_type == "application/pdf":
        prompt = "Extract the complete text contents of this document (which is a resume or job description). Maintain structure where possible."
    else:
        prompt = "Extract all readable text from this file."
        
    result = model.generate_content([part, prompt])
    return result.text

async def generate_application_analysis(payload):
    jd_text = payload.get("jdText")
    resume_text = payload.get("resumeText")
    company_name = payload.get("companyName")
    role_name = payload.get("roleName")
    
    groq_client = get_groq_client()
    
    return await run_resume_analysis_workflow(
        jd_text=jd_text,
        resume_text=resume_text,
        company_name=company_name,
        role_name=role_name,
        groq_client=groq_client
    )

async def generate_analysis_with_gemini(prompt, gemini_client):
    if not gemini_client:
        print("[ai_service] WARNING: Neither Groq nor Gemini keys are set. Returning mock analysis response.")
        if "Interview Evaluator" in prompt:
            return {
                "overallScore": 85,
                "confidenceScore": 80,
                "toneAnalysis": "Mock confident tone.",
                "technicalAccuracy": "Mock technical accuracy.",
                "fillerWords": {"count": 2, "words": ["um", "like"]},
                "recommendations": ["Mock recommendation 1", "Mock recommendation 2", "Mock recommendation 3"],
                "qaReview": [{"question": "Mock Q", "answer": "Mock A", "feedback": "Mock feedback", "score": 85}]
            }
        elif "interview coach" in prompt:
            return {
                "questions": [{"question": "Mock question", "category": "technical", "keyPoints": "Mock key points", "difficulty": "medium"}],
                "tips": "Mock tips."
            }
        elif "Matchmaker" in prompt:
            return {"matches": []}
        elif "Career Connect AI. Write a personalized" in prompt:
            return "This is a mock introduction message."
        elif "Career Twin" in prompt:
            return {
                "readinessScore": 50,
                "readinessReasoning": "Mock reasoning",
                "currentPosition": "Mock position",
                "positionReasoning": "Mock position reasoning",
                "analysis": {
                    "stage": "Entry-Level",
                    "strengths": ["Mock strength"],
                    "criticalWeaknesses": ["Mock weakness"],
                    "missingSkills": [{"skill": "Mock missing skill", "impact": "High"}],
                    "resumeWeaknesses": ["Mock resume weakness"],
                    "competitiveScore": "Low"
                },
                "roadmap": [{
                    "title": "Mock milestone", "whyItMatters": "Mock reason", "estimatedDuration": "2 weeks",
                    "expectedHiringImpact": "High", "deliverables": "Mock deliverables", "completionCriteria": "Mock criteria",
                    "expectedImprovement": "10%"
                }],
                "currentStage": {"stage": "Mock milestone 1", "reasoning": "Mock reason", "confidenceScore": 50},
                "learningResources": [{"title": "Mock resource", "url": "https://example.com", "type": "Mock type"}]
            }
        elif "Career Goal AI Analyst" in prompt:
            return {
                "consistencyScore": 85,
                "analysis": "Mock analysis of goal history.",
                "recommendation": "Continue",
                "careerImpact": "High",
                "nextSteps": ["Mock next step 1", "Mock next step 2"]
            }
        elif "Goal Recommender AI" in prompt:
            return {
                "recommendedGoals": [
                    {
                        "title": "Mock Goal",
                        "category": "Skill Building",
                        "durationDays": 7,
                        "dailyTarget": "Mock daily target",
                        "reasoning": "Mock reasoning",
                        "impact": "High",
                        "difficulty": "Medium",
                        "milestones": ["Mock milestone 1", "Mock milestone 2"]
                    }
                ]
            }
        
        return {
            "company": "Mock Company",
            "role": "Mock Role",
            "signalScore": {
                "fitScore": 75,
                "effort": "Medium",
                "flag": "Green",
                "flagReason": "This is a mock response because API keys are missing or invalid."
            },
            "gaps": [
                {"type": "MISSING KEYWORD", "text": "Mock keyword missing"},
                {"type": "SKILL MISMATCH", "text": "Mock skill mismatch"},
                {"type": "OPPORTUNITY", "text": "Mock opportunity"}
            ],
            "tailoredBullets": [
                {
                    "original": "Original resume bullet point (mock)",
                    "tailored": "Tailored resume bullet point (mock)",
                    "reason": "Mock reason"
                }
            ],
            "outreachMessages": {
                "confident": "Mock confident message.",
                "curious": "Mock curious message.",
                "concise": "Mock concise message."
            },
            "learningResources": [
                {
                    "platform": "YouTube",
                    "title": "Mock learning resource",
                    "link": "https://youtube.com",
                    "description": "Mock description"
                }
            ]
        }
        
    model = gemini_client.GenerativeModel(
        model_name='models/gemini-2.0-flash',
        generation_config={"response_mime_type": "application/json"}
    )
    try:
        result = model.generate_content(prompt)
        return json.loads(clean_json_string(result.text))
    except Exception as e:
        print(f"[ai_service] Gemini API Error: {e}. Falling back to mock response.")
        return await generate_analysis_with_gemini(prompt, None)

async def chat_with_agent(payload):
    message = payload.get("message")
    chat_history = payload.get("chatHistory") or []
    resume_text = payload.get("resumeText")
    current_jd_text = payload.get("currentJdText")
    
    system_prompt = f"""
You are Career Concierge, the personal career agent of the user. You are running inside the "Career Command" AI Job Application Command Centre.
You help candidates improve their resumes, prepare for interviews, write cold outreach emails, and evaluate job fit.
Your personality is professional, encouraging, extremely sharp, and analytical.

Current Candidate Resume:
{resume_text or 'No resume uploaded yet.'}

Current Job Description Under Review:
{current_jd_text or 'No job description under review yet.'}

Instructions:
1. Provide actionable, concise advice.
2. Keep responses brief (under 3 paragraphs) to fit the chat bubble UI.
3. Use bullet points for structural readability.
4. When asked to write or tailor text, focus on highlighting real achievements without fabrication.
"""

    groq_client = get_groq_client()
    gemini_client = get_gemini_client()
    
    messages = [{"role": "system", "content": system_prompt}]
    for h in chat_history:
        messages.append({"role": h.get("role"), "content": h.get("content")})
    messages.append({"role": "user", "content": message})
    
    if groq_client:
        try:
            print("[ai_service] Using Groq API for chat...")
            completion = groq_client.chat.completions.create(
                messages=messages,
                model="llama-3.3-70b-versatile"
            )
            return completion.choices[0].message.content
        except Exception as e:
            print(f"[ai_service] Groq Chat Error, falling back to Gemini: {e}")
            return await chat_with_gemini(messages, gemini_client)
    else:
        print("[ai_service] Using Gemini API for chat...")
        return await chat_with_gemini(messages, gemini_client)

async def chat_with_recruiter_agent(payload):
    message = payload.get("message")
    chat_history = payload.get("chatHistory") or []
    resume_text = payload.get("resumeText")
    current_jd_text = payload.get("currentJdText")
    
    system_prompt = f"""
You are an expert AI Recruiter Simulator conducting a mock interview for the candidate.
Your personality is professional, strict but fair, and observant. You are simulating a real interview environment.

Candidate's Resume:
{resume_text or 'No resume provided.'}

Job Description they are interviewing for:
{current_jd_text or 'No job description provided.'}

Instructions:
1. Speak as a recruiter. Ask ONE question at a time.
2. Based on the candidate's answers, you can ask follow-up questions or move on to the next topic.
3. Your interview should cover:
   - Technical questions (development, CS fundamentals, backend/frontend, scalability, security) based on the JD.
   - Behavioral / "Googliness" questions (teamwork, conflict resolution, leadership).
   - Aptitude / Problem-solving questions.
   - Resume-specific questions (diving deep into hackathons, projects, internships they mentioned).
4. Judge their readiness and confidence implicitly through your responses.
5. Keep your responses short and conversational, as if spoken aloud (no markdown formatting like bolding or bullet points).
6. If the user asks for feedback or to end the interview, provide a comprehensive summary of their performance.
"""

    groq_client = get_groq_client()
    gemini_client = get_gemini_client()
    
    messages = [{"role": "system", "content": system_prompt}]
    for h in chat_history:
        messages.append({"role": h.get("role"), "content": h.get("content")})
    messages.append({"role": "user", "content": message})
    
    if groq_client:
        try:
            print("[ai_service] Using Groq API for recruiter chat...")
            completion = groq_client.chat.completions.create(
                messages=messages,
                model="llama-3.3-70b-versatile"
            )
            return completion.choices[0].message.content
        except Exception as e:
            print(f"[ai_service] Groq Chat Error, falling back to Gemini: {e}")
            return await chat_with_gemini(messages, gemini_client)
    else:
        print("[ai_service] Using Gemini API for recruiter chat...")
        return await chat_with_gemini(messages, gemini_client)

async def evaluate_interview_performance(payload):
    chat_history = payload.get("chatHistory", [])
    jd_text = payload.get("currentJdText", "")
    
    # Format chat history for the prompt
    history_text = ""
    for msg in chat_history:
        role = "Recruiter" if msg.get("role") == "assistant" else "Candidate"
        history_text += f"{role}: {msg.get('content')}\n\n"
        
    prompt = f"""
You are an expert Interview Evaluator. Review the following mock interview transcript between an AI Recruiter and a Candidate.
The interview was for a role described by this Job Description:
{jd_text or 'No job description provided.'}

--- INTERVIEW TRANSCRIPT ---
{history_text}
----------------------------

CRITICAL INSTRUCTIONS FOR SCORING:
- Be brutally honest. Do NOT inflate scores. 
- If the candidate provided no answers, extremely short answers, or nonsensical answers, their scores MUST be 0. 
- If the transcript shows the candidate didn't say anything, overallScore and confidenceScore MUST strictly be 0.

Evaluate the candidate's performance and return a strict JSON object with the following structure:
{{
  "overallScore": <integer 0-100>,
  "confidenceScore": <integer 0-100 based on conciseness and lack of hesitation>,
  "toneAnalysis": "<Analyze the candidate's tone (e.g., nervous, anxious, confident, calm) based on sentence structure, hesitations, and filler words.>",
  "technicalAccuracy": "<Provide a brief assessment of whether the candidate's answers were technically correct or if they made mistakes.>",
  "fillerWords": {{
    "count": <total number of filler words detected>,
    "words": [<array of unique filler words used, e.g., "um", "uh", "like", "you know">]
  }},
  "recommendations": [
    <exactly 3 string bullet points with actionable advice to improve>
  ],
  "qaReview": [
    {{
      "question": "<The recruiter's question>",
      "answer": "<The candidate's exact answer>",
      "feedback": "<1-2 sentences of specific feedback on this answer including whether it was correct>",
      "score": <integer 0-100 for this specific answer>
    }}
  ]
}}

Ensure the output is ONLY valid JSON without any markdown formatting wrappers like ```json.
"""
    
    groq_client = get_groq_client()
    gemini_client = get_gemini_client()
    
    if groq_client:
        try:
            print("[ai_service] Using Groq API for interview evaluation...")
            completion = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.3-70b-versatile",
                response_format={"type": "json_object"}
            )
            resp_text = completion.choices[0].message.content
            return json.loads(clean_json_string(resp_text))
        except Exception as e:
            print(f"[ai_service] Groq Eval Error, falling back to Gemini: {e}")
            return await generate_analysis_with_gemini(prompt, gemini_client)
    else:
        print("[ai_service] Using Gemini API for interview evaluation...")
        return await generate_analysis_with_gemini(prompt, gemini_client)

async def generate_interview_questions(payload):
    target_role = payload.get("targetRole", "")
    resume_text = payload.get("resumeText", "")
    skills = payload.get("skills", "")
    question_types = payload.get("questionTypes", ["technical", "behavioral", "situational"])

    prompt = f"""
You are an expert interview coach. Based on the candidate's profile below, generate a comprehensive set of interview practice questions.

CANDIDATE PROFILE:
Target Role: {target_role or 'Not specified'}
Key Skills: {skills or 'Not specified'}
Resume Summary: {(resume_text[:500] + '...') if len(resume_text) > 500 else resume_text}

Generate questions in these categories: {', '.join(question_types)}.

For each question, provide:
1. The question text
2. The category (technical/behavioral/situational/experience)
3. What the interviewer is looking for (key points to cover)
4. Difficulty level (easy/medium/hard)

Return ONLY a valid JSON object with this structure (no markdown wrapping):
{{
  "questions": [
    {{
      "question": "The interview question",
      "category": "technical",
      "keyPoints": "What the interviewer wants to hear",
      "difficulty": "medium"
    }}
  ],
  "tips": "General interview tips based on the candidate's profile"
}}
"""

    groq_client = get_groq_client()
    gemini_client = get_gemini_client()

    if groq_client:
        try:
            print("[ai_service] Using Groq API for question generation...")
            completion = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.3-70b-versatile",
                response_format={"type": "json_object"}
            )
            return json.loads(clean_json_string(completion.choices[0].message.content))
        except Exception as e:
            print(f"[ai_service] Groq API Error for questions, falling back to Gemini: {e}")
            return await generate_analysis_with_gemini(prompt, gemini_client)
    else:
        print("[ai_service] Using Gemini API for question generation...")
        return await generate_analysis_with_gemini(prompt, gemini_client)

async def chat_with_mentor(payload):
    message = payload.get("message")
    chat_history = payload.get("chatHistory") or []
    resume_text = payload.get("resumeText")
    profile = payload.get("profile", {})
    attachments = payload.get("attachments") or []

    # Process file/image attachments
    attachment_context = ""
    if attachments:
        attachment_context = "\n\n--- ATTACHED FILES ---\n"
        for idx, att in enumerate(attachments):
            try:
                mime = att.get("mime_type", "application/octet-stream")
                data_b64 = att.get("data", "")
                fname = att.get("file_name", f"file_{idx}")
                fbytes = base64.b64decode(data_b64)
                extracted = parse_file_with_gemini(fbytes, mime, fname)
                attachment_context += f"\n[{fname}]:\n{extracted}\n"
            except Exception as e:
                attachment_context += f"\n[{fname}]: (could not parse: {e})\n"
        attachment_context += "\n--- END ATTACHED FILES ---\n"

    # Prepend attachment content to the user message
    if attachment_context:
        message = f"The user has attached the following files. Please use their content to inform your response.\n{attachment_context}\n\nUser message: {message}"

    groq_client = get_groq_client()
    gemini_client = get_gemini_client()

    return await run_mentor_workflow(
        profile=profile,
        resume_text=resume_text,
        chat_history=chat_history,
        message=message,
        groq_client=groq_client,
        gemini_client=gemini_client
    )
async def chat_with_gemini(messages, gemini_client):
    if not gemini_client:
        print("[ai_service] WARNING: No LLM client is available. Set GROQ_API_KEY or GOOGLE_API_KEY.")
        # If this is expected to return JSON (like career twin or interview evaluation), we need to try and parse it or just return a generic text message if it's a chat.
        # But wait, some functions like generate_career_twin use generate_analysis_with_gemini. chat_with_gemini is used by chat agents.
        return "This is a mock AI response. Please configure GROQ_API_KEY or GOOGLE_API_KEY in the backend .env file to enable real AI capabilities."
        
    system_instruction = ""
    for m in messages:
        if m.get("role") == "system":
            system_instruction = m.get("content", "")
            break
            
    model = gemini_client.GenerativeModel(
        model_name='models/gemini-2.0-flash',
        system_instruction=system_instruction
    )
    
    contents = []
    for m in messages:
        if m.get("role") == "system":
            continue
        role = "model" if m.get("role") == "assistant" else "user"
        contents.append({
            "role": role,
            "parts": [{"text": m.get("content", "")}]
        })
        
    history_contents = contents[:-1]
    first_user_index = -1
    for idx, c in enumerate(history_contents):
        if c.get("role") == "user":
            first_user_index = idx
            break
            
    sliced_history = history_contents[first_user_index:] if first_user_index != -1 else []
    
    # start_chat expects list of Content types or structured dicts
    # In python: [{'role': 'user', 'parts': [{'text': 'hi'}]}]
    chat = model.start_chat(history=sliced_history)
    last_msg = contents[-1]["parts"][0]["text"]
    try:
        result = chat.send_message(last_msg)
        return result.text
    except Exception as e:
        print(f"[ai_service] Gemini Chat Error: {e}. Falling back to mock response.")
        return "I am sorry, but my AI connection failed (API Error). Please check that your GROQ_API_KEY or GOOGLE_API_KEY are correct and have quota remaining."

async def generate_network_matches(payload):
    current_profile = payload.get("currentProfile", {})
    candidate_profiles = payload.get("candidateProfiles", [])
    
    if not candidate_profiles:
        return {"matches": []}
        
    prompt = f"""
You are the Career Connect AI Matchmaker. Your job is to find the best peer matches for the user.
Match criteria includes target role, skills, current goals, projects, hackathons, location, and availability.

CURRENT USER:
Target Role: {current_profile.get('targetRole', 'N/A')}
Skills: {current_profile.get('skills', 'N/A')}
Projects: {current_profile.get('projects', 'N/A')}
Hackathons: {current_profile.get('hackathons', 'N/A')}
Goals: {current_profile.get('goals', 'N/A')}
Location: {current_profile.get('location', 'N/A')}
Availability: {current_profile.get('availability', 'N/A')}

CANDIDATES:
{json.dumps(candidate_profiles, indent=2)}

Score each candidate on a scale of 0-100 based on how mutually beneficial a connection would be (e.g. complementary skills, same hackathon, similar goals like Leetcode, study groups).
Return the top 5 matches, sorted by score descending.

Return ONLY a valid JSON object matching this structure:
{{
  "matches": [
    {{
      "email": "candidate email",
      "name": "candidate name",
      "role": "candidate target role",
      "score": 95,
      "reasons": [
        "Same target company",
        "Complementary skills (you have frontend, they have backend)",
        "Both attending HackTX"
      ]
    }}
  ]
}}
"""
    groq_client = get_groq_client()
    gemini_client = get_gemini_client()
    
    if groq_client:
        try:
            print("[ai_service] Using Groq API for network matching...")
            completion = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.3-70b-versatile",
                response_format={"type": "json_object"}
            )
            return json.loads(clean_json_string(completion.choices[0].message.content))
        except Exception as e:
            print(f"[ai_service] Groq Match Error, falling back to Gemini: {e}")
            return await generate_analysis_with_gemini(prompt, gemini_client)
    else:
        print("[ai_service] Using Gemini API for network matching...")
        return await generate_analysis_with_gemini(prompt, gemini_client)

async def generate_ai_introduction(payload):
    current_profile = payload.get("currentProfile", {})
    target_profile = payload.get("targetProfile", {})
    
    prompt = f"""
You are the Career Connect AI. Write a personalized, casual, and friendly introduction message from the Current User to the Target User to initiate networking.
The tone should be natural, not overly formal, and should reference common ground (e.g. shared goals, complementary skills, or location).

CURRENT USER:
Name: {current_profile.get('name', 'N/A')}
Role: {current_profile.get('targetRole', 'N/A')}
Skills: {current_profile.get('skills', 'N/A')}
Goals: {current_profile.get('goals', 'N/A')}

TARGET USER:
Name: {target_profile.get('name', 'N/A')}
Role: {target_profile.get('targetRole', 'N/A')}
Skills: {target_profile.get('skills', 'N/A')}
Goals: {target_profile.get('goals', 'N/A')}

Write only the message text (no JSON, no intro text). Keep it under 4-5 sentences.
"""
    groq_client = get_groq_client()
    gemini_client = get_gemini_client()
    
    if groq_client:
        try:
            print("[ai_service] Using Groq API for intro generation...")
            completion = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.3-70b-versatile"
            )
            return completion.choices[0].message.content
        except Exception as e:
            print(f"[ai_service] Groq Intro Error, falling back to Gemini: {e}")
            return await generate_analysis_with_gemini(prompt, gemini_client)
    else:
        print("[ai_service] Using Gemini API for intro generation...")
        return await generate_analysis_with_gemini(prompt, gemini_client)

async def generate_career_twin(payload):
    resume_text = payload.get("resumeText", "")
    target_role = payload.get("targetRole", "")
    skills = payload.get("skills", "")
    projects = payload.get("projects", "")
    experience = payload.get("experience", "")
    education = payload.get("education", "")
    completed_goals = payload.get("completedGoals", [])
    
    goals_context = ""
    if completed_goals:
        goals_context = "\nRECENTLY COMPLETED GOALS BY CANDIDATE:\n"
        for g in completed_goals:
            goals_context += f"- {g.get('title')} (Category: {g.get('category')})\n"
        goals_context += "CRITICAL INSTRUCTION: Since the candidate has recently completed these goals, DO NOT list them as missing skills if they address a skill gap. Instead, explicitly increase their readiness score and acknowledge this progress in the analysis.\n"

    prompt = f"""
You are the Career Twin AI. Your task is to analyze the candidate's profile against realistic market expectations for their target role.

CANDIDATE DATA:
Resume: {resume_text}
Skills: {skills}
Experience: {experience}
Projects: {projects}
Education: {education}
Target Role: {target_role}
{goals_context}

--------------------------------------------------
ANALYSIS
--------------------------------------------------
First determine the candidate's current career stage (e.g. Beginner, Entry-Level, Junior, Mid-Level, Senior).
Then compare the candidate against realistic market expectations for the requested role. Compare against candidates who are actually getting hired today.

Identify:
1. Core strengths (only genuine strengths supported by evidence).
2. Critical weaknesses (issues that will significantly reduce interview chances).
3. Missing technical skills (rank by impact).
4. Resume weaknesses (e.g. Weak projects, No measurable achievements, Poor ATS optimization, Lack of production experience).
5. Competitive Score (Estimate how competitive the candidate is today).

Return: Current Readiness: XX% and explain exactly why.

--------------------------------------------------
CURRENT POSITION
--------------------------------------------------
Determine where the candidate currently stands (e.g. Learning Fundamentals, Building Projects, Job Ready, Interview Ready, Industry Competitive, Top Tier Candidate). Provide reasoning.

--------------------------------------------------
ROADMAP
--------------------------------------------------
Generate a realistic roadmap toward the target role.
Each milestone should include: Title, Why it matters, Estimated duration, Expected hiring impact, Deliverables, Completion criteria, and Expected improvement in hiring probability (e.g. from 32% to 45%).
Roadmaps must be practical. Avoid vague tasks.

--------------------------------------------------
CURRENT POSITION MARKER
--------------------------------------------------
Determine which roadmap milestone the candidate is currently on.
Return: Current Stage, Reasoning, Confidence Score.

--------------------------------------------------
RESOURCE RECOMMENDATIONS
--------------------------------------------------
Recommend only high-quality learning resources (Official documentation, Roadmap.sh, FreeCodeCamp, etc.).
Prioritize the top 20% of improvements that would produce roughly 80% of the increase in interview and hiring probability.

--------------------------------------------------
OUTPUT FORMAT
--------------------------------------------------
Return ONLY a valid JSON object strictly matching this format (no markdown code blocks):
{{
  "readinessScore": 32,
  "readinessReasoning": "...",
  "currentPosition": "Building Projects",
  "positionReasoning": "...",
  "analysis": {{
    "stage": "Entry-Level",
    "strengths": ["..."],
    "criticalWeaknesses": ["..."],
    "missingSkills": [
      {{"skill": "Skill Name", "impact": "High"}}
    ],
    "resumeWeaknesses": ["..."],
    "competitiveScore": "Low"
  }},
  "roadmap": [
    {{
      "title": "Milestone Title",
      "whyItMatters": "...",
      "estimatedDuration": "2 weeks",
      "expectedHiringImpact": "High",
      "deliverables": "...",
      "completionCriteria": "...",
      "expectedImprovement": "45%"
    }}
  ],
  "currentStage": {{
    "stage": "Milestone 1",
    "reasoning": "...",
    "confidenceScore": 85
  }},
  "learningResources": [
    {{
      "title": "Resource Title",
      "url": "Search URL or actual URL",
      "type": "Official documentation"
    }}
  ]
}}
"""

    groq_client = get_groq_client()
    gemini_client = get_gemini_client()

    if groq_client:
        try:
            print("[ai_service] Using Groq API for career twin...")
            completion = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.3-70b-versatile",
                response_format={"type": "json_object"}
            )
            return json.loads(clean_json_string(completion.choices[0].message.content))
        except Exception as e:
            print(f"[ai_service] Groq Twin Error, falling back to Gemini: {e}")
            return await generate_analysis_with_gemini(prompt, gemini_client)
    else:
        print("[ai_service] Using Gemini API for career twin...")
        return await generate_analysis_with_gemini(prompt, gemini_client)

async def analyze_goal_performance(payload):
    goal = payload.get("goal", {})
    history = payload.get("history", {}) # { "YYYY-MM-DD": "Completed" | "Missed" | "Skipped" }
    
    prompt = f"""
You are the Career Goal AI Analyst. 
The candidate just ended a goal. Review their daily check-in history for this goal and provide a strict, data-driven performance report.

GOAL DETAILS:
Title: {goal.get('title')}
Category: {goal.get('category')}
Duration (Days): {goal.get('durationDays')}
Daily Target: {goal.get('dailyTarget') or 'None'}

CHECK-IN HISTORY:
{json.dumps(history, indent=2)}

Calculate their consistency, identify strongest/weakest periods, and estimate the career impact of this effort. 
Provide practical recommendations for their next step (Continue, Extend, or Replace this goal).

Return ONLY a valid JSON object matching this structure:
{{
  "consistency": "Good",
  "completionRate": 75,
  "strongestPeriod": "Early in the week",
  "weakestPeriod": "Weekends",
  "estimatedSkillImprovement": "Moderate improvement in core concepts.",
  "recommendations": ["Recommendation 1", "Recommendation 2"],
  "nextStep": "Extend"
}}
"""
    groq_client = get_groq_client()
    gemini_client = get_gemini_client()

    if groq_client:
        try:
            print("[ai_service] Using Groq API for goal analysis...")
            completion = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.3-70b-versatile",
                response_format={"type": "json_object"}
            )
            return json.loads(clean_json_string(completion.choices[0].message.content))
        except Exception as e:
            print(f"[ai_service] Groq Error, falling back to Gemini: {e}")
            return await generate_analysis_with_gemini(prompt, gemini_client)
    else:
        print("[ai_service] Using Gemini API for goal analysis...")
        return await generate_analysis_with_gemini(prompt, gemini_client)

async def generate_goal_recommendations(payload):
    target_role = payload.get("targetRole", "")
    skills = payload.get("skills", "")
    missing_skills = payload.get("missingSkills", [])
    
    prompt = f"""
You are the Career Action AI. Based on the candidate's target role and their identified missing skills, recommend 5 extremely practical, high-impact goals they should add to their daily tracker to boost their hiring readiness.

Target Role: {target_role}
Current Skills: {skills}
Missing Skills Identified by Career Twin: {json.dumps(missing_skills)}

Each goal must be highly specific and actionable (e.g., "Build a Dockerized REST API" rather than "Learn Docker").

Return ONLY a valid JSON object matching this structure:
{{
  "recommendations": [
    {{
      "title": "Goal Title",
      "description": "Short description of why it matters.",
      "category": "Projects",
      "durationDays": 14,
      "dailyTarget": "Code for 1 hour"
    }}
  ]
}}
"""
    groq_client = get_groq_client()
    gemini_client = get_gemini_client()

    if groq_client:
        try:
            print("[ai_service] Using Groq API for goal recommendations...")
            completion = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.3-70b-versatile",
                response_format={"type": "json_object"}
            )
            return json.loads(clean_json_string(completion.choices[0].message.content))
        except Exception as e:
            print(f"[ai_service] Groq Error, falling back to Gemini: {e}")
            return await generate_analysis_with_gemini(prompt, gemini_client)
    else:
        print("[ai_service] Using Gemini API for goal recommendations...")
        return await generate_analysis_with_gemini(prompt, gemini_client)
